"""
Assessment DOCX Generator — matches Streamlit's exact format.
Generates Questions and Answer docs for each assessment type.
"""

import sys
import json
import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


TYPE_LONG_NAMES = {
    'WA (SAQ)': 'Written Assessment – Short Answer Questions (SAQ)',
    'WA-SAQ': 'Written Assessment – Short Answer Questions (SAQ)',
    'WA(Q&A)': 'Written Assessment – Short Answer Questions (SAQ)',
    'PP': 'Practical Performance (PP)',
    'CS': 'Case Study (CS)',
    'PRJ': 'Project (PRJ)',
    'ASGN': 'Assignment (ASGN)',
    'OI': 'Oral Interview (OI)',
    'DEM': 'Demonstration (DEM)',
    'RP': 'Role Play (RP)',
    'OQ': 'Oral Questioning (OQ)',
}

TYPE_SHORT_NAMES = {
    'WA (SAQ)': 'Written Assessment (SAQ)',
    'WA-SAQ': 'Written Assessment (SAQ)',
    'PP': 'Practical Performance',
    'CS': 'Case Study',
    'PRJ': 'Project',
    'ASGN': 'Assignment',
    'OI': 'Oral Interview',
    'DEM': 'Demonstration',
    'RP': 'Role Play',
    'OQ': 'Oral Questioning',
}


def _set_cell_border(cell, sz='4', color='000000'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for b in ['top', 'left', 'bottom', 'right']:
        e = OxmlElement(f'w:{b}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), sz)
        e.set(qn('w:color'), color)
        tcBorders.append(e)
    tcPr.append(tcBorders)


def _build_ref_string(q):
    refs = []
    k_id = q.get('knowledge_id', '')
    a_ids = q.get('ability_id', [])
    if isinstance(a_ids, str):
        a_ids = [a_ids] if a_ids else []
    if k_id:
        refs.append(k_id)
    refs.extend([a for a in a_ids if a])
    return f" ({', '.join(refs)})" if refs else ''


def _format_duration(duration):
    if not duration:
        return ''
    duration = str(duration).strip()
    if 'min' in duration.lower() and 'hour' not in duration.lower() and 'hr' not in duration.lower():
        return duration
    total_mins = 0
    hr_match = re.search(r'(\d+\.?\d*)\s*(?:hour|hr)', duration, re.IGNORECASE)
    min_match = re.search(r'(\d+)\s*min', duration, re.IGNORECASE)
    if hr_match:
        total_mins += int(float(hr_match.group(1)) * 60)
    if min_match:
        total_mins += int(min_match.group(1))
    return f'{total_mins} mins' if total_mins > 0 else duration


def _add_line(doc, text, size=12, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    return p


def build_question_doc(assessment, course_title, output_path):
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    a_type = assessment.get('type', assessment.get('code', ''))
    code = assessment.get('code', a_type)
    duration = _format_duration(assessment.get('duration', ''))
    questions = assessment.get('questions', [])
    short_name = TYPE_SHORT_NAMES.get(a_type, TYPE_SHORT_NAMES.get(code, a_type))
    long_name = TYPE_LONG_NAMES.get(a_type, TYPE_LONG_NAMES.get(code, a_type))
    num_q = len(questions)

    # Title
    _add_line(doc, course_title, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_line(doc, short_name, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Section A: Trainee Information
    _add_line(doc, 'A: Trainee Information:', bold=True, size=12)
    _add_line(doc, 'Trainee Name (as Per NRIC): ______________________________', size=12)
    _add_line(doc, 'Last three digits and alphabet of NRIC/FIN: _______________', size=12)
    _add_line(doc, 'Date: _______________', size=12)

    # Section B: Assessment Instruction
    _add_line(doc, 'B: Assessment Instruction', bold=True, size=12)
    _add_line(doc, f'This is the {long_name}', size=12)
    if duration:
        _add_line(doc, f'Duration: {duration}', size=12)
    _add_line(doc, f'1. The assessor will pass the questions in hard copy to you. There are {num_q} questions. You need to answer all the questions.', size=12)
    _add_line(doc, '2. This is an open-book exam that must be completed individually.', size=12)
    _add_line(doc, '3. You need to get all answers correct to be competent.', size=12)
    _add_line(doc, 'Submission Procedure:', bold=True, size=12)
    _add_line(doc, '1. Please pass the hard copy to the assessor after completion.', size=12)

    # Section C: Questions
    _add_line(doc, 'C: Questions and Answers', bold=True, size=12)

    for i, q in enumerate(questions, 1):
        scenario = q.get('scenario', '')
        question_text = q.get('question_statement', q.get('question', ''))
        ref = _build_ref_string(q)

        # Scenario (plain paragraph)
        if scenario:
            _add_line(doc, scenario, size=12)

        # Question
        _add_line(doc, f'Q{i}. {question_text}{ref}', bold=True, size=12)

        # Answer: label
        _add_line(doc, 'Answer:', bold=True, size=12)

        # Empty answer box (bordered table, 1 cell, multiple empty rows)
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        _set_cell_border(cell)
        cell.paragraphs[0].text = ''
        for _ in range(8):
            cell.add_paragraph('')

    # For Official Use Only section (at end, matching Streamlit)
    # Horizontal line
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    _add_line(doc, 'For Official Use Only', bold=True, size=12)
    _add_line(doc, 'Grade: __________(C / NYC)', size=12)
    _add_line(doc, 'Assessor Name: _______________    Assessor NRIC: _______________', size=12)
    _add_line(doc, 'Date: _______________    Signature: _______________', size=12)

    doc.save(output_path)
    return output_path


def build_answer_doc(assessment, course_title, output_path):
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    a_type = assessment.get('type', assessment.get('code', ''))
    code = assessment.get('code', a_type)
    questions = assessment.get('questions', [])
    short_name = TYPE_SHORT_NAMES.get(a_type, TYPE_SHORT_NAMES.get(code, a_type))
    long_name = TYPE_LONG_NAMES.get(a_type, TYPE_LONG_NAMES.get(code, a_type))

    # Title
    _add_line(doc, f'Answers to {course_title}', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_line(doc, short_name, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Horizontal separator
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    for i, q in enumerate(questions, 1):
        scenario = q.get('scenario', '')
        question_text = q.get('question_statement', q.get('question', ''))
        ref = _build_ref_string(q)
        answers = q.get('answer', [])

        if scenario:
            _add_line(doc, scenario, italic=True, size=12)

        _add_line(doc, f'Q{i}. {question_text}{ref}', bold=True, size=12)

        # Answer box
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        _set_cell_border(cell)
        cell.paragraphs[0].text = ''
        p = cell.paragraphs[0]
        run = p.add_run('Suggestive answers (not exhaustive):')
        run.bold = True
        run.italic = True
        run.font.size = Pt(12)

        if isinstance(answers, list):
            for ans in answers:
                ap = cell.add_paragraph()
                ap.paragraph_format.left_indent = Cm(0.5)
                r = ap.add_run(f'• {ans}')
                r.font.size = Pt(12)
        elif answers:
            ap = cell.add_paragraph()
            r = ap.add_run(str(answers))
            r.font.size = Pt(12)

    doc.save(output_path)
    return output_path


def sanitize_filename(name):
    return re.sub(r'[^\w\-. ()]', '', name)[:80].strip()


def main():
    if len(sys.argv) != 3:
        print(json.dumps({'error': 'Usage: generate-assessment.py <input_json> <output_dir>'}))
        sys.exit(1)

    input_path = sys.argv[1]
    output_dir = sys.argv[2]

    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        course_title = data.get('course_title', 'Course')
        assessments = data.get('assessments', [])

        if not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)

        clean_title = sanitize_filename(course_title)
        generated = []
        for assessment in assessments:
            a_type = assessment.get('type', assessment.get('code', 'Assessment'))
            type_safe = sanitize_filename(a_type)

            # Match Streamlit filename format: "WA (SAQ) - Python Programming for Finance.docx"
            q_path = os.path.join(output_dir, f'{type_safe} - {clean_title}.docx')
            a_path = os.path.join(output_dir, f'Answer to {type_safe} - {clean_title}.docx')

            build_question_doc(assessment, course_title, q_path)
            build_answer_doc(assessment, course_title, a_path)

            generated.append({
                'type': a_type,
                'question_path': q_path,
                'answer_path': a_path,
            })

        print(json.dumps({'success': True, 'generated': generated}))
    except Exception as e:
        import traceback
        print(json.dumps({'error': str(e), 'traceback': traceback.format_exc()}))
        sys.exit(1)


if __name__ == '__main__':
    main()
