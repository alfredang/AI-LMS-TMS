"""
Lesson Plan Generator — Pure Python barrier scheduling algorithm.
Matches Streamlit's timetable_generator.py exactly.

Usage: python generate-lp.py <context_json_path> <output_path>
"""

import sys
import json
import os
import re
import math
from datetime import datetime
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm, Inches, Pt
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(SCRIPT_DIR)
TEMPLATE_DIR = os.path.join(PROJECT_DIR, 'public', 'templates', 'courseware')

# ─── Constants (matching Streamlit) ───
DAY_START = 9 * 60          # 9:00 AM in minutes
DAY_END = 18 * 60           # 6:00 PM in minutes
LUNCH_START = 12 * 60 + 30  # 12:30 PM
LUNCH_DURATION = 45          # minutes
LUNCH_END = LUNCH_START + LUNCH_DURATION  # 1:15 PM
ASSESS_START = 16 * 60       # 4:00 PM
MIN_SESSION = 15             # Minimum session before barrier


def _parse_hours(val):
    """Parse hours from various formats: '16 hrs', '16', 16, '16 hours 30 minutes'"""
    if isinstance(val, (int, float)):
        return float(val)
    if not val:
        return 0
    val = str(val).strip()
    # Try "X hours Y minutes" format
    h_match = re.search(r'(\d+)\s*(?:hour|hr)', val, re.IGNORECASE)
    m_match = re.search(r'(\d+)\s*(?:min)', val, re.IGNORECASE)
    if h_match:
        hours = int(h_match.group(1))
        mins = int(m_match.group(1)) if m_match else 0
        return hours + mins / 60
    # Try plain number
    num_match = re.search(r'(\d+\.?\d*)', val)
    return float(num_match.group(1)) if num_match else 0


def _round5(mins):
    """Round minutes to nearest 5."""
    return int(5 * round(mins / 5))


def _fmt_time(mins):
    """Format minutes since midnight to '9:00 AM' format."""
    h = int(mins // 60)
    m = int(mins % 60)
    ampm = 'AM' if h < 12 else 'PM'
    h12 = h if h <= 12 else h - 12
    if h12 == 0:
        h12 = 12
    return f"{h12}:{m:02d} {ampm}"


def build_schedule(context):
    """Build lesson plan schedule using barrier algorithm."""
    total_hours = _parse_hours(context.get('Total_Course_Duration_Hours', '16'))
    instr_hours = _parse_hours(context.get('Total_Training_Hours', ''))
    assess_hours = _parse_hours(context.get('Total_Assessment_Hours', '0'))

    # Try to extract correct hours from extractedText (parsed CP) if available
    extracted_text = context.get('extractedText', '')
    if extracted_text:
        import re as _re
        cr_m = _re.search(r'Classroom\s*\w*\s*\|\s*([\d.]+)\s*hours?', extracted_text, _re.IGNORECASE)
        pr_m = _re.search(r'Practical\s*.*?Practicum\w*\s*\|\s*([\d.]+)\s*hours?', extracted_text, _re.IGNORECASE)
        el_m = _re.search(r'E-Learning\s*\|\s*([\d.]+)\s*hours?', extracted_text, _re.IGNORECASE)
        as_m = _re.search(r'\|\s*Assessment\s*\|\s*([\d.]+)\s*hours?', extracted_text, _re.IGNORECASE)
        td_m = _re.search(r'Total\s*Duration\s*\|\s*([\d.]+)\s*hours?', extracted_text, _re.IGNORECASE)

        calc_instr = 0
        if cr_m: calc_instr += float(cr_m.group(1))
        if pr_m: calc_instr += float(pr_m.group(1))
        if el_m: calc_instr += float(el_m.group(1))

        if calc_instr > 0:
            instr_hours = calc_instr
        if as_m:
            assess_hours = float(as_m.group(1))
        if td_m:
            total_hours = float(td_m.group(1))

        # Also count topics from raw CP text (T1:, T2:, etc.)
        topic_count_from_cp = len(_re.findall(r'T\d+\s*:', extracted_text))
        print(json.dumps({'debug_hours': True, 'classroom': cr_m.group(1) if cr_m else None, 'practical': pr_m.group(1) if pr_m else None, 'assess': as_m.group(1) if as_m else None, 'total': td_m.group(1) if td_m else None, 'calc_instr': calc_instr, 'topics_from_cp': topic_count_from_cp}), file=sys.stderr)

    if not instr_hours:
        instr_hours = total_hours - assess_hours

    num_days = max(1, round(total_hours / 8)) if total_hours >= 8 else 1

    # Collect individual topic blocks with LU grouping (matching Streamlit — each topic = separate row)
    lus = context.get('Learning_Units', [])
    topic_blocks = []  # Each entry = one topic
    for i, lu in enumerate(lus):
        topics = lu.get('Topics', [])
        methods = lu.get('Instructional_Methods', [])
        methods_str = ', '.join(methods) if methods else '-'
        lu_title = lu.get('LU_Title', f'Learning Unit {i+1}')

        if not topics:
            topic_blocks.append({
                'lu_num': i + 1, 'lu_title': lu_title,
                'topic_desc': lu_title, 'methods': methods_str,
                'is_first_in_lu': True,
            })
        else:
            for j, t in enumerate(topics):
                title = t.get('Topic_Title', t.get('title', f'Topic {j+1}'))
                topic_blocks.append({
                    'lu_num': i + 1, 'lu_title': lu_title,
                    'topic_desc': f"T{j+1}: {title}",
                    'methods': methods_str,
                    'is_first_in_lu': (j == 0),
                })

    if not topic_blocks:
        return {'num_days': num_days, 'days': {}, 'error': 'No Learning Units found'}

    total_topics = len(topic_blocks)
    per_topic = (instr_hours * 60) / total_topics if total_topics > 0 else 60

    # Schedule each topic individually with barrier algorithm
    days = {}
    t_idx = 0
    t_remaining = per_topic
    is_contd = False
    last_lu_num = -1

    for day in range(1, num_days + 1):
        slots = []
        current = DAY_START
        is_last_day = (day == num_days)
        lunch_done = False
        assess_mins = int(assess_hours * 60)
        instr_end = (DAY_END - assess_mins) if (is_last_day and assess_mins > 0) else DAY_END

        while t_idx < len(topic_blocks) and current < instr_end:
            block = topic_blocks[t_idx]

            # LU header row when entering new LU
            if block['lu_num'] != last_lu_num:
                contd = " (Cont'd)" if is_contd else ""
                slots.append({
                    'timing': '', 'duration': '', 'methods': '',
                    'description': f"LU{block['lu_num']}: {block['lu_title']}{contd}",
                    'is_lu_header': True,
                })
                last_lu_num = block['lu_num']

            # Lunch barrier
            if not lunch_done and current >= LUNCH_START:
                slots.append({
                    'timing': f"{_fmt_time(current)} - {_fmt_time(LUNCH_END)}",
                    'duration': f"{LUNCH_DURATION} mins",
                    'description': 'Lunch Break', 'methods': '-',
                })
                current = LUNCH_END
                lunch_done = True
                # Re-insert LU header after lunch if continuing same LU
                if is_contd:
                    slots.append({
                        'timing': '', 'duration': '', 'methods': '',
                        'description': f"LU{block['lu_num']}: {block['lu_title']} (Cont'd)",
                        'is_lu_header': True,
                    })
                continue

            next_barrier = LUNCH_START if not lunch_done else instr_end
            available = next_barrier - current
            if available <= 0:
                break

            contd_suffix = " (Cont'd)" if is_contd else ""
            desc = f"{block['topic_desc']}{contd_suffix}"

            if t_remaining <= available:
                end = _round5(current + t_remaining)
                end = min(end, instr_end)
                dur = end - current
                if dur > 0:
                    slots.append({
                        'timing': f"{_fmt_time(current)} - {_fmt_time(end)}",
                        'duration': f"{int(dur)} mins",
                        'description': desc, 'methods': block['methods'],
                    })
                current = end
                t_idx += 1
                t_remaining = per_topic
                is_contd = False
            elif available >= MIN_SESSION:
                slots.append({
                    'timing': f"{_fmt_time(current)} - {_fmt_time(next_barrier)}",
                    'duration': f"{int(available)} mins",
                    'description': desc, 'methods': block['methods'],
                })
                t_remaining -= available
                is_contd = True
                current = next_barrier
            else:
                if not lunch_done:
                    slots.append({
                        'timing': f"{_fmt_time(current)} - {_fmt_time(LUNCH_END)}",
                        'duration': f"{LUNCH_DURATION} mins",
                        'description': 'Lunch Break', 'methods': '-',
                    })
                    current = LUNCH_END
                    lunch_done = True
                else:
                    slots.append({
                        'timing': f"{_fmt_time(current)} - {_fmt_time(next_barrier)}",
                        'duration': f"{int(available)} mins",
                        'description': 'Break', 'methods': '-',
                    })
                    current = next_barrier

        if not lunch_done:
            if current < LUNCH_START:
                slots.append({
                    'timing': f"{_fmt_time(current)} - {_fmt_time(LUNCH_START)}",
                    'duration': f"{int(LUNCH_START - current)} mins",
                    'description': 'Break', 'methods': '-',
                })
                current = LUNCH_START
            slots.append({
                'timing': f"{_fmt_time(current)} - {_fmt_time(LUNCH_END)}",
                'duration': f"{LUNCH_DURATION} mins",
                'description': 'Lunch Break', 'methods': '-',
            })
            current = LUNCH_END

        if is_last_day and assess_mins > 0:
            am_details = context.get('Assessment_Methods_Details', [])
            am_parts = [am.get('Assessment_Method', '') for am in am_details if am.get('Assessment_Method')]
            am_desc = "Assessment: " + ", ".join(am_parts) if am_parts else f"Assessment ({assess_mins} mins)"

            am_start = max(current, DAY_END - assess_mins)
            if current < am_start:
                slots.append({
                    'timing': f"{_fmt_time(current)} - {_fmt_time(am_start)}",
                    'duration': f"{int(am_start - current)} mins",
                    'description': 'Break', 'methods': '-',
                })
                current = am_start
            am_end = min(current + assess_mins, DAY_END)
            slots.append({
                'timing': f"{_fmt_time(current)} - {_fmt_time(am_end)}",
                'duration': f"{int(am_end - current)} mins",
                'description': am_desc, 'methods': 'Assessment',
            })
            current = am_end

        if current < DAY_END:
            slots.append({
                'timing': f"{_fmt_time(current)} - {_fmt_time(DAY_END)}",
                'duration': f"{int(DAY_END - current)} mins",
                'description': 'Break', 'methods': '-',
            })

        days[day] = slots

    return {
        'num_days': num_days,
        'instructional_hours': instr_hours,
        'assessment_hours': assess_hours,
        'per_topic_mins': round(per_topic, 1),
        'days': days,
    }


def generate_lp_docx(context, schedule, output_path):
    """Generate Lesson Plan DOCX from template + schedule data."""
    # First render the cover page template
    tpl_path = os.path.join(TEMPLATE_DIR, 'LP_template_v2.docx')
    if not os.path.exists(tpl_path):
        tpl_path = os.path.join(TEMPLATE_DIR, 'LP_TGS-Ref-No_Course-Title_v1.docx')

    now = datetime.now()
    tpl = DocxTemplate(tpl_path)

    tpl_ctx = {
        'Course_Title': context.get('Course_Title', ''),
        'TGS_Ref_No': context.get('TGS_Ref_No', ''),
        'Name_of_Organisation': context.get('Name_of_Organisation', ''),
        'UEN': context.get('UEN', '') or '201200696W',
        'Date': now.strftime("%d %b %Y"),
        'Year': str(now.year),
    }

    # Add logo
    logo_path = os.path.join(TEMPLATE_DIR, 'tertiary_logo.png')
    if os.path.exists(logo_path):
        try:
            tpl_ctx['company_logo'] = InlineImage(tpl, logo_path, width=Mm(50))
        except:
            tpl_ctx['company_logo'] = ''
    else:
        tpl_ctx['company_logo'] = ''

    tpl.render(tpl_ctx, autoescape=True)
    tpl.save(output_path)

    # Now open the saved doc and add the schedule tables
    doc = Document(output_path)

    # Add title
    p = doc.add_paragraph()
    run = p.add_run(f"Lesson Plan: {context.get('Course_Title', '')}")
    run.bold = True
    run.font.size = Pt(14)

    # Add metadata
    num_days = schedule['num_days']
    instr_hrs = schedule['instructional_hours']
    assess_hrs = schedule['assessment_hours']
    all_methods = set()
    for lu in context.get('Learning_Units', []):
        for m in lu.get('Instructional_Methods', []):
            all_methods.add(m)

    meta = doc.add_paragraph()
    meta.add_run(f"Course Duration: {num_days} Day(s) (9:00 AM - 6:00 PM daily)\n").font.size = Pt(10)
    meta.add_run(f"Total Training Hours: {instr_hrs} hrs\n").font.size = Pt(10)
    meta.add_run(f"Total Assessment Hours: {assess_hrs} hrs\n").font.size = Pt(10)
    meta.add_run(f"Instructional Methods: {', '.join(sorted(all_methods)) or 'N/A'}\n").font.size = Pt(10)
    meta.add_run(f"Organisation: {context.get('Name_of_Organisation', '')}").font.size = Pt(10)

    # Add day-by-day tables
    for day_num in sorted(schedule['days'].keys()):
        slots = schedule['days'][day_num]

        doc.add_paragraph()
        day_heading = doc.add_paragraph()
        run = day_heading.add_run(f"Day {day_num}")
        run.bold = True
        run.font.size = Pt(12)

        # Create 4-column table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        headers = ['Timing', 'Duration', 'Description', 'Instructional Methods']
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)

        # Data rows
        for slot in slots:
            if slot.get('is_lu_header'):
                # LU header row — merge all cells
                row = table.add_row()
                row.cells[0].merge(row.cells[3])
                row.cells[0].text = slot.get('description', '')
                for paragraph in row.cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(10)
            else:
                row = table.add_row()
                row.cells[0].text = slot.get('timing', '')
                row.cells[1].text = slot.get('duration', '')
                row.cells[2].text = slot.get('description', '')
                row.cells[3].text = slot.get('methods', '')

                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

    doc.save(output_path)
    return output_path


def main():
    if len(sys.argv) != 3:
        print(json.dumps({'error': 'Usage: generate-lp.py <context_json_path> <output_path>'}))
        sys.exit(1)

    context_path = sys.argv[1]
    output_path = sys.argv[2]

    try:
        with open(context_path, 'r', encoding='utf-8') as f:
            context = json.load(f)

        schedule = build_schedule(context)

        if schedule.get('error'):
            print(json.dumps({'error': schedule['error']}))
            sys.exit(1)

        result_path = generate_lp_docx(context, schedule, output_path)
        # Convert days keys to strings for JSON serialization
        schedule_out = {**schedule, 'days': {str(k): v for k, v in schedule['days'].items()}}
        print(json.dumps({'success': True, 'path': result_path, 'schedule': schedule_out}))
    except Exception as e:
        print(json.dumps({'error': str(e)}))
        sys.exit(1)


if __name__ == '__main__':
    main()
