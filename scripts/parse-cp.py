"""
CP Document Parser — matches Streamlit's parse_cp_document() exactly.
Uses python-docx to extract paragraphs and tables in document order.

Usage: python parse-cp.py <input_file_path>
Outputs: parsed text to stdout as JSON
"""

import sys
import json
import os
import re
import tempfile
import base64


def parse_docx(file_path):
    """Parse DOCX file extracting paragraphs and tables in document order.
    Matches Streamlit's parse_cp_document() logic."""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(file_path)
    body = doc.element.body
    text_parts = []

    for element in body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

        if tag == 'p':
            # Paragraph
            para_text = element.text or ''
            # Get all runs
            for run in element.findall(qn('w:r')):
                t = run.find(qn('w:t'))
                if t is not None and t.text:
                    para_text = para_text if para_text else ''

            # Use the paragraph object for proper text extraction
            for para in doc.paragraphs:
                if para._element is element:
                    para_text = para.text
                    # Preserve heading styles as markdown
                    if para.style and 'Heading' in para.style.name:
                        level = ''.join(filter(str.isdigit, para.style.name)) or '1'
                        prefix = '#' * min(int(level), 4)
                        para_text = f"{prefix} {para_text}"
                    break

            if para_text.strip():
                text_parts.append(para_text)

        elif tag == 'tbl':
            # Table — convert to markdown format
            for table in doc.tables:
                if table._element is element:
                    rows_text = []
                    prev_row = None
                    for row in table.rows:
                        cells = []
                        for cell in row.cells:
                            cell_text = cell.text.strip().replace('\n', ' ')
                            cells.append(cell_text)
                        # Deduplicate merged cells
                        row_str = ' | '.join(cells)
                        if row_str != prev_row:
                            rows_text.append(f"| {row_str} |")
                            prev_row = row_str

                    if rows_text:
                        # Add header separator after first row
                        if len(rows_text) > 1:
                            num_cols = rows_text[0].count('|') - 1
                            separator = '|' + '|'.join(['---'] * num_cols) + '|'
                            rows_text.insert(1, separator)
                        text_parts.append('\n'.join(rows_text))
                    break

    return '\n\n'.join(text_parts)


def parse_xlsx(file_path):
    """Parse XLSX file extracting all sheets as pipe-separated data.
    Matches Streamlit's parse_cp_document() logic for Excel."""
    import openpyxl

    wb = openpyxl.load_workbook(file_path, data_only=True)
    all_text = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sheet_rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(cell).strip() if cell is not None else '' for cell in row]
            if any(cells):
                sheet_rows.append(' | '.join(cells))
        if sheet_rows:
            all_text.append(f"## {sheet_name}\n" + '\n'.join(sheet_rows))

    return '\n\n'.join(all_text)


def trim_docx_content(text):
    """Trim to relevant CP sections: Part 1 through Part 4."""
    start_match = re.search(r'Part\s*1[\s\S]*?Particulars\s*of\s*Course', text, re.IGNORECASE)
    end_match = re.search(r'Part\s*\d+[\s\S]*?Facilities\s*and\s*Resources', text, re.IGNORECASE)
    if start_match:
        start_idx = start_match.start()
        if end_match:
            text = text[start_idx:end_match.end()]
        else:
            text = text[start_idx:]
    return re.sub(r'\n{3,}', '\n\n', text).strip()


def trim_xlsx_content(text):
    """Trim to relevant sections."""
    start_idx = text.find('1 - Course Particulars')
    end_idx = text.find('4 - Declarations')
    if start_idx >= 0:
        text = text[start_idx:end_idx] if end_idx >= 0 else text[start_idx:]
    return re.sub(r'\n{3,}', '\n\n', text).strip()


def main():
    if len(sys.argv) < 2:
        print(json.dumps({'error': 'Usage: parse-cp.py <file_path> [is_base64]'}))
        sys.exit(1)

    file_path = sys.argv[1]
    is_base64 = len(sys.argv) > 2 and sys.argv[2] == 'base64'

    try:
        actual_path = file_path

        if is_base64:
            # Read base64 content and write to temp file
            with open(file_path, 'r') as f:
                b64_content = f.read().strip()
            # Detect file type from original filename (passed as arg 3)
            orig_name = sys.argv[3] if len(sys.argv) > 3 else 'file.docx'
            ext = os.path.splitext(orig_name)[1].lower()
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
            tmp.write(base64.b64decode(b64_content))
            tmp.close()
            actual_path = tmp.name

        if actual_path.lower().endswith('.docx'):
            text = parse_docx(actual_path)
            text = trim_docx_content(text)
        elif actual_path.lower().endswith(('.xlsx', '.xls')):
            text = parse_xlsx(actual_path)
            text = trim_xlsx_content(text)
        else:
            with open(actual_path, 'r', encoding='utf-8') as f:
                text = f.read()

        # Limit size
        if len(text) > 80000:
            text = text[:80000] + '\n[Content truncated]'

        print(json.dumps({'success': True, 'text': text}))

        # Cleanup temp file
        if is_base64:
            os.unlink(actual_path)

    except Exception as e:
        print(json.dumps({'error': str(e)}))
        sys.exit(1)


if __name__ == '__main__':
    main()
